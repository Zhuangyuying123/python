from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
def set_paragraph_format(paragraph, font_size=Pt(12), line_spacing=1.5, first_line_indent=Pt(24)):
    for run in paragraph.runs:
        run.font.size = font_size
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.first_line_indent = first_line_indent


doc = Document('【排版前】基于Python爬虫的全北京小区信息文本分析.docx')

# 获取文档的所有段落
paragraphs = doc.paragraphs


# 遍历所有段落
for paragraph in doc.paragraphs:
    for index, paragraph in enumerate(doc.paragraphs):
        set_paragraph_format(paragraph)  # 设置默认段落格式

        # 第二步：设置第一行的格式
        if index == 0:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(16)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)

        # 第三步和第四步：设置以一、二、三....等开头的段落的格式
        elif paragraph.text.startswith(('一', '二', '三', '四', '五')):
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(14)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.first_line_indent = Pt(0)

        # 第三步和第四步：设置以（一）', '（二）', '（三）','（四）','（五）等开头的段落的格式
        elif paragraph.text.startswith(('（一）', '（二）', '（三）', '（四）', '（五）')):
            for run in paragraph.runs:
                run.font.size = Pt(13)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.first_line_indent = Pt(0)

        # 第五步：设置以图开头的段落的格式
        elif paragraph.text.startswith('图'):
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
        elif paragraph.text.startswith('图'):
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.first_line_indent = Pt(0)

for para in doc.paragraphs :
    txt = para.text# 段落中不包含任何文本
    if not txt:
        contentId = (para.part.inline_shapes[0]._inline.graphic.
        graphicData.pic.blipFill.blip.embed)
        contentType = para.part.related_parts[contentId].content_type
        # 段落中包含图片
        if contentType.startswith('image'):
            # 段落对齐方式设置为居中
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 设置默认字体为 'Times New Roman'
doc.styles['Normal'].font.name = 'Times New Roman'

# 为每个段落的 runs 设置中文字体
for run in paragraph.runs:
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')




# 保存文档
doc.save('【排版后】基于Python爬虫的全北京小区信息文本分析.docx')
